# import os
# import json
# from typing import Dict, Any, List
# import os
# import tempfile
# from docx import Document
# from docx2pdf import convert
# from azure.core.credentials import AzureKeyCredential
# from azure.ai.documentintelligence import DocumentIntelligenceClient
# from azure.ai.documentintelligence.models import AnalyzeDocumentRequest


# def get_polygon_bbox(polygon):
#     if not polygon or len(polygon) != 8:
#         return None
#     xs = polygon[0::2]
#     ys = polygon[1::2]
#     return [min(xs), min(ys), max(xs), max(ys)]



# def extract_layout_to_structured_json(
#     file_bytes: bytes,
#     source_name: str
# ) -> Dict[str, Any]:
#     """
#     Extract document using Azure Document Intelligence.
#     - Handles .docx by converting to PDF first
#     - Uses prebuilt-layout for best layout/table accuracy
#     """
#     client = DocumentIntelligenceClient(
#         endpoint=os.getenv("DOC_INTELLIGENCE_ENDPOINT"),
#         credential=AzureKeyCredential(os.getenv("DOC_INTELLIGENCE_KEY"))
#     )

#     ext = os.path.splitext(source_name)[1].lower()

#     # -------------------------------------------------
#     # Handle DOCX: convert to PDF first
#     # -------------------------------------------------
#     if ext == ".docx":
#         print(f"Converting DOCX to PDF: {source_name}")
#         # Save bytes to temp .docx file
#         with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
#             tmp_docx.write(file_bytes)
#             tmp_docx_path = tmp_docx.name

#         # Convert to PDF
#         tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
#         convert(tmp_docx_path, tmp_pdf_path)

#         # Read PDF bytes
#         with open(tmp_pdf_path, "rb") as f:
#             file_bytes = f.read()

#         # Clean up temp files
#         os.unlink(tmp_docx_path)
#         os.unlink(tmp_pdf_path)

#         source_name = source_name.replace(".docx", ".pdf")  # update name

#     # -------------------------------------------------
#     # Now process (PDF or image bytes) with prebuilt-layout
#     # -------------------------------------------------
#     print(f"Using prebuilt-layout for: {source_name}")

#     request = AnalyzeDocumentRequest(bytes_source=file_bytes)
#     poller = client.begin_analyze_document("prebuilt-layout", body=request)
#     result = poller.result()

#     structured_doc = {
#         "document_name": source_name,
#         "model": "prebuilt-layout",
#         "pages": []
#     }

#     # Helper: get bbox
#     def get_bbox(polygon):
#         if not polygon or len(polygon) != 8:
#             return None
#         xs = polygon[0::2]
#         ys = polygon[1::2]
#         return [min(xs), min(ys), max(xs), max(ys)]

#     # -------------------------------------------------
#     # Process paragraphs
#     # -------------------------------------------------
#     page_map = {}

#     for para in result.paragraphs or []:
#         if not para.content or not para.bounding_regions:
#             continue

#         region = para.bounding_regions[0]
#         page_number = region.page_number
#         bbox = get_bbox(region.polygon)

#         if page_number not in page_map:
#             page_map[page_number] = {"page_number": page_number, "blocks": []}

#         page_map[page_number]["blocks"].append({
#             "block_id": f"p{page_number}_para_{len(page_map[page_number]['blocks'])}",
#             "block_type": "paragraph",
#             "text": para.content.strip(),
#             "bbox": bbox
#         })

#     # -------------------------------------------------
#     # Process tables (unchanged)
#     # -------------------------------------------------
#     previous_table_page = None
#     previous_headers = None

#     for t_idx, table in enumerate(result.tables or []):
#         if not table.bounding_regions:
#             continue

#         table_page = table.bounding_regions[0].page_number
#         table_bbox = get_bbox(table.bounding_regions[0].polygon)

#         is_continuation = previous_table_page is not None and table_page == previous_table_page + 1

#         grid = [[""] * table.column_count for _ in range(table.row_count)]
#         for cell in table.cells:
#             r, c = cell.row_index, cell.column_index
#             if 0 <= r < len(grid) and 0 <= c < len(grid[r]):
#                 grid[r][c] = cell.content or ""

#         headers = []
#         rows = []

#         header_row_indices = {cell.row_index for cell in table.cells if cell.kind == "columnHeader"}

#         if header_row_indices and not is_continuation:
#             header_row = min(header_row_indices)
#             headers = grid[header_row]
#             rows = [row for i, row in enumerate(grid) if i != header_row]
#             previous_headers = headers
#         elif is_continuation and previous_headers:
#             headers = previous_headers
#             rows = grid
#         else:
#             headers = []
#             rows = grid

#         if headers:
#             rows = [
#                 row for row in rows
#                 if any(cell.strip() != hdr.strip() for cell, hdr in zip(row, headers))
#             ]

#         table_block = {
#             "block_id": f"table_{t_idx}",
#             "block_type": "table",
#             "page_number": table_page,
#             "bbox": table_bbox,
#             "headers": headers,
#             "rows": rows
#         }

#         if table_page not in page_map:
#             page_map[table_page] = {"page_number": table_page, "blocks": []}
#         page_map[table_page]["blocks"].append(table_block)

#         previous_table_page = table_page

#     # -------------------------------------------------
#     # Sort blocks per page (top-to-bottom)
#     # -------------------------------------------------
#     for page_num in sorted(page_map.keys()):
#         page = page_map[page_num]
#         page["blocks"].sort(key=lambda b: b["bbox"][1] if b["bbox"] else 0)
#         structured_doc["pages"].append(page)


#     # -------------------------------------------------
#     # Save raw output for debugging
#     # -------------------------------------------------
#     debug_file = "ICH_layout_structured.json"
#     with open(debug_file, "w", encoding="utf-8") as f:
#         json.dump(structured_doc, f, ensure_ascii=False, indent=2)
#     print(f"Structured layout saved to: {debug_file}")

#     print("✅ Layout JSON saved to: ich_layout_structured.json")
#     return structured_doc










############################################ New Code like source ########################################

import os
import json
from typing import Dict, Any, List
import tempfile
from io import BytesIO
from docx import Document
from docx2pdf import convert
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import (
    AnalyzeDocumentRequest,
    DocumentContentFormat
)


def get_polygon_bbox(polygon):
    if not polygon or len(polygon) != 8:
        return None
    xs = polygon[0::2]
    ys = polygon[1::2]
    return [min(xs), min(ys), max(xs), max(ys)]


def extract_layout_to_structured_json(
    file_bytes: bytes,
    source_name: str
) -> Dict[str, Any]:
    """
    Extract document using Azure Document Intelligence.
    - Converts .docx to PDF first (safe handling)
    - Uses prebuilt-layout
    - Enables high resolution OCR
    - Uses markdown output
    - Intelligent table continuation detection
    """

    client = DocumentIntelligenceClient(
        endpoint=os.getenv("DOC_INTELLIGENCE_ENDPOINT"),
        credential=AzureKeyCredential(os.getenv("DOC_INTELLIGENCE_KEY"))
    )

    ext = os.path.splitext(source_name)[1].lower()

    # -------------------------------------------------
    # Handle DOCX safely
    # -------------------------------------------------
    if ext == ".docx":
        print(f"Converting DOCX to PDF: {source_name}")

        tmp_docx_path = None
        tmp_pdf_path = None

        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
                tmp_docx.write(file_bytes)
                tmp_docx_path = tmp_docx.name

            tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
            convert(tmp_docx_path, tmp_pdf_path)

            with open(tmp_pdf_path, "rb") as f:
                file_bytes = f.read()

            source_name = source_name.replace(".docx", ".pdf")
            print("DOCX conversion successful")

        except Exception as e:
            print(f"DOCX → PDF conversion failed: {e}")
            doc = Document(BytesIO(file_bytes))
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

            return {
                "document_name": source_name,
                "model": "fallback-text",
                "pages": [{
                    "page_number": 1,
                    "blocks": [{
                        "block_id": "fallback",
                        "block_type": "paragraph",
                        "text": full_text,
                        "bbox": None
                    }]
                }]
            }

        finally:
            for path in [tmp_docx_path, tmp_pdf_path]:
                if path and os.path.exists(path):
                    try:
                        os.remove(path)
                    except Exception:
                        pass

    # -------------------------------------------------
    # Analyze using prebuilt-layout (Improved settings)
    # -------------------------------------------------
    print(f"Using prebuilt-layout for: {source_name}")

    request = AnalyzeDocumentRequest(bytes_source=file_bytes)

    poller = client.begin_analyze_document(
        "prebuilt-layout",
        body=request,
        features=["ocrHighResolution"],
        output_content_format=DocumentContentFormat.MARKDOWN
    )

    result = poller.result(timeout=600)  # 10 minutes max

    structured_doc = {
        "document_name": source_name,
        "model": "prebuilt-layout",
        "pages": []
    }

    page_map = {}

    # -------------------------------------------------
    # Process paragraphs
    # -------------------------------------------------
    for para in result.paragraphs or []:
        if not para.content or not para.bounding_regions:
            continue

        region = para.bounding_regions[0]
        page_number = region.page_number
        bbox = get_polygon_bbox(region.polygon)

        if page_number not in page_map:
            page_map[page_number] = {"page_number": page_number, "blocks": []}

        page_map[page_number]["blocks"].append({
            "block_id": f"p{page_number}_para_{len(page_map[page_number]['blocks'])}",
            "block_type": "paragraph",
            "text": para.content.strip(),
            "bbox": bbox
        })

    # -------------------------------------------------
    # Process tables (Intelligent continuation logic)
    # -------------------------------------------------
    previous_table_page = None
    previous_headers = None

    for t_idx, table in enumerate(result.tables or []):
        if not table.bounding_regions:
            continue

        table_page = table.bounding_regions[0].page_number
        table_bbox = get_polygon_bbox(table.bounding_regions[0].polygon)

        # Build grid first
        grid = [[""] * table.column_count for _ in range(table.row_count)]
        for cell in table.cells:
            r, c = cell.row_index, cell.column_index
            if 0 <= r < len(grid) and 0 <= c < len(grid[r]):
                grid[r][c] = cell.content or ""

        is_continuation = False

        # Smart continuation detection
        if previous_table_page is not None and table_page == previous_table_page + 1:
            if previous_headers:
                header_row_indices = {
                    cell.row_index for cell in table.cells if cell.kind == "columnHeader"
                }

                if header_row_indices:
                    current_header_row = min(header_row_indices)
                    current_headers = grid[current_header_row]

                    if len(current_headers) == len(previous_headers):
                        similarity = sum(
                            1 for a, b in zip(current_headers, previous_headers)
                            if a.strip().lower().split("(")[0]
                            == b.strip().lower().split("(")[0]
                        )

                        if similarity >= len(previous_headers) * 0.7:
                            is_continuation = True

        headers = []
        rows = []

        header_row_indices = {
            cell.row_index for cell in table.cells if cell.kind == "columnHeader"
        }

        if header_row_indices and not is_continuation:
            header_row = min(header_row_indices)
            headers = grid[header_row]
            rows = [row for i, row in enumerate(grid) if i != header_row]
            previous_headers = headers

        elif is_continuation and previous_headers:
            headers = previous_headers
            rows = grid

        else:
            headers = []
            rows = grid

        # Remove duplicate header rows
        if headers:
            rows = [
                row for row in rows
                if any(cell.strip() != hdr.strip()
                       for cell, hdr in zip(row, headers))
            ]

        table_block = {
            "block_id": f"table_{t_idx}",
            "block_type": "table",
            "page_number": table_page,
            "bbox": table_bbox,
            "headers": headers,
            "rows": rows
        }

        if table_page not in page_map:
            page_map[table_page] = {"page_number": table_page, "blocks": []}

        page_map[table_page]["blocks"].append(table_block)
        previous_table_page = table_page

    # -------------------------------------------------
    # Sort blocks top-to-bottom
    # -------------------------------------------------
    for page_num in sorted(page_map.keys()):
        page = page_map[page_num]
        page["blocks"].sort(key=lambda b: b["bbox"][1] if b["bbox"] else 0)
        structured_doc["pages"].append(page)

    # -------------------------------------------------
    # Save debug output
    # -------------------------------------------------
    debug_file = "ICH_layout_structured.json"
    with open(debug_file, "w", encoding="utf-8") as f:
        json.dump(structured_doc, f, ensure_ascii=False, indent=2)

    print(f"Structured layout saved to: {debug_file}")
    print("✅ Layout JSON saved to: ich_layout_structured.json")

    return structured_doc