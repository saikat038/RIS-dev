"""
PDF → Word converter using pdfplumber.
Preserves text, tables, and basic formatting with high accuracy.

Install dependencies:
    pip install pdfplumber python-docx tqdm

For Word → PDF (commented-out function), also install:
    pip install docx2pdf
"""

import os
import logging
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# CORE HELPERS
# ─────────────────────────────────────────────

def _add_table_to_doc(doc: Document, table_data: list[list]) -> None:
    """Write a pdfplumber table into a python-docx Table."""
    if not table_data:
        return

    # Filter out fully-None rows
    rows = [r for r in table_data if any(cell is not None for cell in r)]
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < col_count:
                text = str(cell_text) if cell_text is not None else ""
                tbl.rows[r_idx].cells[c_idx].text = text

    doc.add_paragraph()  # breathing room after table


def _extract_page_content(page) -> tuple[list, list]:
    """
    Returns (text_blocks, tables) for a single pdfplumber page.
    text_blocks: list of str
    tables:      list of list[list]  (one per table found)
    """
    tables = page.extract_tables()
    table_bboxes = [t.bbox for t in page.find_tables()] if tables else []

    # Crop out table regions so we don't double-extract their text
    cropped = page
    for bbox in table_bboxes:
        try:
            cropped = cropped.filter(
                lambda obj, bb=bbox: not (
                    bb[0] <= obj.get("x0", 0) <= bb[2]
                    and bb[1] <= obj.get("top", 0) <= bb[3]
                )
            )
        except Exception:
            pass  # If filtering fails, continue with full page text

    raw_text = cropped.extract_text(x_tolerance=2, y_tolerance=3) or ""
    text_blocks = [line for line in raw_text.splitlines() if line.strip()]

    return text_blocks, tables


# ─────────────────────────────────────────────
# MAIN CONVERSION FUNCTION
# ─────────────────────────────────────────────

def pdf_to_word(pdf_path: str | Path, output_path: str | Path) -> bool:
    """
    Convert a single PDF to a .docx Word document.

    Args:
        pdf_path:    Path to source .pdf file.
        output_path: Path to write the .docx output.

    Returns:
        True on success, False on failure.
    """
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc = Document()

        # ── Document-level style tweaks ──────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)

            for page_num, page in enumerate(
                tqdm(pdf.pages, desc=f"  Pages ({pdf_path.name})", leave=False, unit="pg"),
                start=1,
            ):
                # ── Page separator ───────────────────────────
                if page_num > 1:
                    doc.add_page_break()

                text_blocks, tables = _extract_page_content(page)

                # ── Write text blocks ────────────────────────
                for line in text_blocks:
                    para = doc.add_paragraph(line)
                    para.paragraph_format.space_after = Pt(2)

                # ── Write tables ─────────────────────────────
                for table_data in tables:
                    _add_table_to_doc(doc, table_data)

        doc.save(output_path)
        return True

    except Exception as exc:
        log.error("Failed to convert '%s': %s", pdf_path.name, exc)
        return False


# ─────────────────────────────────────────────
# BATCH PROCESSING
# ─────────────────────────────────────────────

def convert_folder(
    pdf_folder: str | Path,
    output_folder: str | Path | None = None,
) -> dict:
    """
    Convert all PDFs in a folder to Word documents.

    Args:
        pdf_folder:    Folder containing .pdf files (searched recursively).
        output_folder: Destination folder. Defaults to <pdf_folder>/word_output/

    Returns:
        Summary dict with keys: total, success, failed, skipped, failed_files.
    """
    pdf_folder = Path(pdf_folder)
    if not pdf_folder.exists():
        raise FileNotFoundError(f"Folder not found: {pdf_folder}")

    if output_folder is None:
        output_folder = pdf_folder / "word_output"
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    pdf_files = sorted(pdf_folder.rglob("*.pdf"))

    if not pdf_files:
        log.warning("No PDF files found in '%s'", pdf_folder)
        return {"total": 0, "success": 0, "failed": 0, "skipped": 0, "failed_files": []}

    log.info("Found %d PDF(s) → output: %s", len(pdf_files), output_folder)

    results = {"total": len(pdf_files), "success": 0, "failed": 0, "skipped": 0, "failed_files": []}

    with tqdm(pdf_files, desc="Converting PDFs", unit="file") as pbar:
        for pdf_path in pbar:
            pbar.set_postfix(file=pdf_path.name[:30])

            # Maintain subfolder structure inside output_folder
            relative = pdf_path.relative_to(pdf_folder)
            docx_path = output_folder / relative.with_suffix(".docx")

            # Skip if already converted
            if docx_path.exists():
                log.debug("Skipping (already exists): %s", docx_path.name)
                results["skipped"] += 1
                continue

            ok = pdf_to_word(pdf_path, docx_path)
            if ok:
                results["success"] += 1
            else:
                results["failed"] += 1
                results["failed_files"].append(str(pdf_path))

    # ── Summary ──────────────────────────────────────────────
    log.info(
        "Done. ✓ %d converted | ✗ %d failed | ⟳ %d skipped",
        results["success"], results["failed"], results["skipped"],
    )
    if results["failed_files"]:
        log.warning("Failed files:\n  %s", "\n  ".join(results["failed_files"]))

    return results


# ─────────────────────────────────────────────
# WORD → PDF  (highest quality)
# Uncomment when needed. Requires: pip install docx2pdf
# On Linux you also need LibreOffice installed:
#   sudo apt install libreoffice
# ─────────────────────────────────────────────

# def word_to_pdf_hq(
#     word_folder: str | Path,
#     output_folder: str | Path | None = None,
# ) -> dict:
#     """
#     Convert all .docx files in a folder to high-quality PDFs.
#
#     Uses docx2pdf which calls:
#       - Microsoft Word on Windows/macOS (highest fidelity)
#       - LibreOffice on Linux (good fidelity)
#
#     Args:
#         word_folder:   Folder containing .docx files.
#         output_folder: Destination folder. Defaults to <word_folder>/pdf_hq_output/
#
#     Returns:
#         Summary dict with keys: total, success, failed, failed_files.
#     """
#     from docx2pdf import convert  # lazy import — only needed for this function
#
#     word_folder = Path(word_folder)
#     if output_folder is None:
#         output_folder = word_folder / "pdf_hq_output"
#     output_folder = Path(output_folder)
#     output_folder.mkdir(parents=True, exist_ok=True)
#
#     docx_files = sorted(word_folder.rglob("*.docx"))
#     if not docx_files:
#         log.warning("No .docx files found in '%s'", word_folder)
#         return {"total": 0, "success": 0, "failed": 0, "failed_files": []}
#
#     log.info("Found %d Word file(s) → output: %s", len(docx_files), output_folder)
#     results = {"total": len(docx_files), "success": 0, "failed": 0, "failed_files": []}
#
#     with tqdm(docx_files, desc="Word → PDF (HQ)", unit="file") as pbar:
#         for docx_path in pbar:
#             pbar.set_postfix(file=docx_path.name[:30])
#             relative = docx_path.relative_to(word_folder)
#             pdf_out = output_folder / relative.with_suffix(".pdf")
#             pdf_out.parent.mkdir(parents=True, exist_ok=True)
#             try:
#                 convert(str(docx_path), str(pdf_out))
#                 results["success"] += 1
#             except Exception as exc:
#                 log.error("Failed '%s': %s", docx_path.name, exc)
#                 results["failed"] += 1
#                 results["failed_files"].append(str(docx_path))
#
#     log.info(
#         "Done. ✓ %d converted | ✗ %d failed",
#         results["success"], results["failed"],
#     )
#     return results


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    # 👇 DEFAULT PATHS (used if no CLI args given)
    DEFAULT_PDF_FOLDER = r"C:\Users\SaikatSome\Downloads\miriyala\OCU400-Protocol, SAP & Interim CSR\actual files\pdfs"
    DEFAULT_OUTPUT_FOLDER = r"C:\Users\SaikatSome\Downloads\miriyala\OCU400-Protocol, SAP & Interim CSR\actual files\output_docs"

    parser = argparse.ArgumentParser(
        description="Batch convert PDFs to Word documents."
    )

    parser.add_argument(
        "pdf_folder",
        nargs="?",  # 👈 makes it optional
        default=DEFAULT_PDF_FOLDER,
        help="Path to folder containing PDF files"
    )

    parser.add_argument(
        "--output", "-o",
        default=DEFAULT_OUTPUT_FOLDER,
        help="Output folder for .docx files"
    )

    args = parser.parse_args()

    print("📂 Input Folder :", args.pdf_folder)
    print("📁 Output Folder:", args.output)

    convert_folder(
        pdf_folder=args.pdf_folder,
        output_folder=args.output,
    )