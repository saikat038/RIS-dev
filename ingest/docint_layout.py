import os
import json
import os
import re
import tempfile
from io import BytesIO
from docx import Document
from docx2pdf import convert
from typing import Dict, Any, List
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
# from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.ai.documentintelligence.models import DocumentContentFormat
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest


def get_polygon_bbox(polygon):
    if not polygon or len(polygon) != 8:
        return None
    xs = polygon[0::2]
    ys = polygon[1::2]
    return [min(xs), min(ys), max(xs), max(ys)]


def merge_empty_header_columns(headers, rows):
    """
    Merge columns where header is empty with the previous column.
    """

    i = 0
    while i < len(headers):

        if headers[i].strip() == "" and i > 0:

            # merge header
            headers[i-1] = headers[i-1].strip()

            # merge row values
            for row in rows:
                if i < len(row):
                    if row[i].strip():
                        row[i-1] = (row[i-1] + " " + row[i]).strip()

            # remove the column
            headers.pop(i)
            for row in rows:
                if i < len(row):
                    row.pop(i)

            continue

        i += 1

    return headers, rows



def normalize_cohort_headers(headers):
    """
    Fix headers like:
    'Cohort (N=2) 2' -> 'Cohort 2 (N=2)'
    """

    fixed = []

    for h in headers:
        h = re.sub(r"\s+", " ", h).strip()

        match = re.match(r"^(.*)\(N=(\d+)\)\s*(\d+)$", h)
        if match:
            prefix = match.group(1).strip()
            n_value = match.group(2)
            cohort_num = match.group(3)

            h = f"{prefix} {cohort_num} (N={n_value})"

        fixed.append(h)

    return fixed


# def merge_numeric_header_fragments(headers, rows):
#     """
#     Merge columns where the header is just a number (e.g., '2')
#     with the previous column header.
#     """

#     i = 0
#     while i < len(headers):

#         header = headers[i].strip()

#         if re.fullmatch(r"\d+", header) and i > 0:

#             # merge header text
#             prev = headers[i-1].replace("\n", " ").strip()
#             headers[i-1] = f"{prev} {header}"

#             # merge row values
#             for row in rows:
#                 if i < len(row):
#                     row[i-1] = (row[i-1] + " " + row[i]).strip()

#             # remove this column
#             headers.pop(i)

#             for row in rows:
#                 if i < len(row):
#                     row.pop(i)

#             continue

#         i += 1

#     return headers, rows


def fix_section_number(text: str, previous_heading: str | None):
    """
    Fix OCR errors where section numbers start with 0.
    Uses previous heading prefix if available.
    """

    if not text:
        return text

    # Only check headings starting with 0.
    if not re.match(r"^0\.\d+", text):
        return text

    # If we have previous heading context
    if previous_heading:
        match = re.match(r"^(\d+)\.", previous_heading)
        if match:
            prefix = match.group(1)
            return re.sub(r"^0\.", f"{prefix}.", text)

    return text


def extract_layout_to_structured_json(
    file_bytes: bytes,
    source_name: str
) -> Dict[str, Any]:
    """
    Extract document using Azure Document Intelligence.
    - Converts .docx to PDF first (required for good quality)
    - Uses prebuilt-layout for layout + table accuracy
    """
    client = DocumentIntelligenceClient(
        endpoint=os.getenv("DOC_INTELLIGENCE_ENDPOINT"),
        credential=AzureKeyCredential(os.getenv("DOC_INTELLIGENCE_KEY")),
        api_version= "2024-11-30"
    )
    # client = DocumentAnalysisClient(
    #     endpoint=os.getenv("DOC_INTELLIGENCE_ENDPOINT"),
    #     credential=AzureKeyCredential(os.getenv("DOC_INTELLIGENCE_KEY"))
    # )

    ext = os.path.splitext(source_name)[1].lower()

    # ────────────────────────────────────────────────
    # Handle DOCX: convert to PDF
    # ────────────────────────────────────────────────
    if ext == ".docx":
        print(f"Converting DOCX to PDF: {source_name}")
        try:
            # Save to temp file (docx2pdf needs real files on disk)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
                tmp_in.write(file_bytes)
                tmp_in_path = tmp_in.name

            tmp_out_path = tmp_in_path.replace(".docx", ".pdf")
            convert(tmp_in_path, tmp_out_path)

            with open(tmp_out_path, "rb") as f:
                file_bytes = f.read()

            # Clean up
            os.unlink(tmp_in_path)
            os.unlink(tmp_out_path)

            source_name = source_name.replace(".docx", ".pdf")  # update name
            print("Conversion successful")
        except Exception as e:
            print(f"DOCX → PDF conversion failed: {e}")
            # Fallback: extract basic text with python-docx
            doc = Document(BytesIO(file_bytes))
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {
                "document_name": source_name,
                "model": "fallback-text",
                "pages": [{
                    "page_number": 1,
                    "blocks": [{
                        "block_id": "fallback",
                        "block_type": "paragraph",
                        "text": full_text,
                        "bbox": None
                    }]
                }]
            }

    # ────────────────────────────────────────────────
    # Now process bytes (PDF or image) with prebuilt-layout
    # ────────────────────────────────────────────────
    print(f"Analyzing with prebuilt-layout: {source_name}")
    
    request = AnalyzeDocumentRequest(bytes_source=file_bytes)
    # poller = client.begin_analyze_document("prebuilt-layout", body=request, features=["ocrHighResolution"], output_content_format=DocumentContentFormat.MARKDOWN )
    poller = client.begin_analyze_document("prebuilt-layout", body=request, output_content_format=DocumentContentFormat.MARKDOWN )
    # poller = client.begin_analyze_document(
    #     "prebuilt-layout",
    #     document=file_bytes
    # )
    result = poller.result()
    with open("full_result.json", "w", encoding="utf-8") as f:
        json.dump(result.as_dict(), f, ensure_ascii=False, indent=2)
    
    # for page in result.pages:
    #     for word in page.words:
    #         if abs(word.polygon[1] - 5.67) < 0.1:  # near the same y position
    #             print(word.content, word.polygon)
    structured_doc = {
        "document_name": source_name,
        "model": "prebuilt-layout",
        "pages": []
    }

    previous_heading = None

    # Helper: get bbox
    def get_bbox(polygon):
        if not polygon or len(polygon) != 8:
            return None
        xs = polygon[0::2]
        ys = polygon[1::2]
        return [min(xs), min(ys), max(xs), max(ys)]

    # Process paragraphs (full paragraphs from model)
    page_map = {}

    table_regions = []

    for table in result.tables or []:

        for region in table.bounding_regions or []:

            bbox = get_bbox(region.polygon)

            if bbox:
                table_regions.append({
                    "page": region.page_number,
                    "bbox": bbox
                })

    for para in result.paragraphs or []:
        role = getattr(para, "role", None)
        if not para.content or not para.bounding_regions:
            continue

        region = para.bounding_regions[0]
        page_number = region.page_number
        bbox = get_bbox(region.polygon)
        # skip paragraphs that belong to tables
        inside_table = False

        for tb in table_regions:

            if tb["page"] != page_number:
                continue

            inter = (
                max(0, min(bbox[2], tb["bbox"][2]) - max(bbox[0], tb["bbox"][0]))
                *
                max(0, min(bbox[3], tb["bbox"][3]) - max(bbox[1], tb["bbox"][1]))
            )

            para_area = (
                (bbox[2] - bbox[0]) *
                (bbox[3] - bbox[1])
            )

            if para_area > 0 and (inter / para_area) > 0.7:
                inside_table = True
                break

        if inside_table:
            continue

        if page_number not in page_map:
            page_map[page_number] = {"page_number": page_number, "blocks": []}

        text = para.content.strip()

        # fix OCR section numbering
        text = fix_section_number(text, previous_heading)

        # detect heading pattern
        if re.match(r"^\d+(\.\d+)+\.?\s", text):
            previous_heading = text

        role = getattr(para, "role", None)

        page_map[page_number]["blocks"].append({
            "block_id": f"p{page_number}_para_{len(page_map[page_number]['blocks'])}",
            "block_type": "paragraph",
            "text": text,
            "bbox": bbox,

            # NEW
            "azure_role": role,
            "is_heading": role == "sectionHeading"
        })

    # Process tables (your existing logic is fine)
    previous_table_page = None
    previous_headers = None

    for t_idx, table in enumerate(result.tables or []):
        if not table.bounding_regions:
            continue
        # Skip very small metadata tables
        if table.row_count <= 2 and table.column_count <= 3:
            continue

        table_page = table.bounding_regions[0].page_number
        table_bbox = get_bbox(table.bounding_regions[0].polygon)

        is_continuation = False

        max_row = max(cell.row_index for cell in table.cells) + 1
        max_col = max(cell.column_index for cell in table.cells) + 1

        grid = [[""] * max_col for _ in range(max_row)]

        for cell in table.cells:

            r = cell.row_index
            c = cell.column_index

            content = (cell.content or "").strip()

            if (
                0 <= r < max_row
                and 0 <= c < max_col
            ):

                # don't overwrite already populated cell
                if not grid[r][c]:
                    grid[r][c] = content

                        

        if previous_table_page is not None and table_page == previous_table_page + 1:

            if previous_headers:

                current_header_indices = {
                    cell.row_index
                    for cell in table.cells
                    if cell.kind == "columnHeader"
                }

                if current_header_indices:
                    current_header_row = min(current_header_indices)
                    current_headers = grid[current_header_row]
                else:
                    current_headers = grid[0]

                # continuation check must ALWAYS run
                if len(current_headers) == len(previous_headers):

                    similarity = sum(
                        1 for a, b in zip(current_headers, previous_headers)
                        if a.strip().lower().split("(")[0]
                        == b.strip().lower().split("(")[0]
                    )

                    if similarity >= len(previous_headers) * 0.7:
                        is_continuation = True


        

        headers = []
        rows = []

        header_cells = [
            cell for cell in table.cells
            if cell.kind == "columnHeader"
        ]

        if header_cells and not is_continuation:

            # ONLY use top-most header row
            top_header_row = min(
                cell.row_index for cell in header_cells
            )

            top_headers = [
                cell for cell in header_cells
                if cell.row_index == top_header_row
            ]

            # sort left-to-right
            top_headers = sorted(
                top_headers,
                key=lambda x: x.column_index
            )

            headers = [
                (cell.content or "").strip()
                for cell in top_headers
            ]

            header_row_indices = {
                cell.row_index for cell in header_cells
            }

            rows = [
                row for i, row in enumerate(grid)
                if i not in header_row_indices
            ]

            previous_headers = headers

        elif is_continuation and previous_headers:

            headers = previous_headers
            rows = grid

        else:

            headers = []
            rows = grid

        if headers:
            rows = [
                row for row in rows
                if len(row) != len(headers) or any(cell.strip() != hdr.strip() for cell, hdr in zip(row, headers))
            ]
        
        # apply cleanup only if headers look broken

        # if any(h.strip() == "" for h in headers):
        #     headers, rows = merge_empty_header_columns(headers, rows)

        headers = normalize_cohort_headers(headers)

        table_block = {
            "block_id": f"table_{t_idx}",
            "block_type": "table",
            "page_number": table_page,
            "bbox": table_bbox,
            "headers": headers,
            "rows": rows
        }

        if table_page not in page_map:
            page_map[table_page] = {"page_number": table_page, "blocks": []}
        page_map[table_page]["blocks"].append(table_block)

        previous_table_page = table_page
        if not headers:
            previous_headers = None

    # Sort blocks per page top-to-bottom
    for page_num in sorted(page_map.keys()):
        page = page_map[page_num]
        page["blocks"].sort(key=lambda b: (b["bbox"][1], b["bbox"][0]) if b["bbox"] else (0,0))
        structured_doc["pages"].append(page)

    # -------------------------------------------------
    # Save raw structured layout
    # -------------------------------------------------
    # Create folder if it doesn't exist
    output_dir = "OCR output/OCR raw"
    os.makedirs(output_dir, exist_ok=True)

    # Build output file path using doc_id
    output_json = os.path.join(output_dir, f"{source_name}_raw_structured.json")

    # Save JSON
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(structured_doc, f, ensure_ascii=False, indent=2)

            

    # with open("sample_layout_structured.json", "w", encoding="utf-8") as f:
    #     json.dump(structured_doc, f, ensure_ascii=False, indent=2)

    print("✅ Layout JSON saved to: sample_layout_structured.json")
    return structured_doc