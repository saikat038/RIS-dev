import re
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ==========================================================
# 1️⃣ RAW TABLE INPUT (Your Exact Table)
# ==========================================================

RAW_TABLE = """
Characteristic Statistic/Category | Cohort 1 (N=4) | Cohort 2 (N=3) | Cohort 3 (N=3) | Phase 2 (N=8) | MTD (N=11) | Total (N=18)
|---|---|---|---|---|---|---|
Genetic Diagnosis, n (%) |  |  |  |  |  |
NR2E3 Mutation | 0 ( 0.0) | 2 ( 66.7) | 2 ( 66.7) | 4 ( 50.0) | 6 ( 54.5) | 8 ( 44.4)
RHO Mutation | 4 (100.0) | 1 ( 33.3) | 1 ( 33.3) | 4 ( 50.0) | 5 ( 45.5) | 10 ( 55.6)
Subgroup - Mutation Subtype, n (%) |  |  |  |  |  |
Subgroup 1 - Biallelic autosomal recessive NR2E3 mutations | 0 ( 0.0) | 2 ( 66.7) | 1 ( 33.3) | 2 ( 25.0) | 3 ( 27.3) | 5 ( 27.8)
Subgroup 2 - Autosomal dominant NR2E3 mutations | 0 ( 0.0) | 0 ( 0.0) | 1 ( 33.3) | 2 ( 25.0) | 3 ( 27.3) | 3 ( 16.7)
Subgroup 3 - Autosomal dominant RHO mutations | 4 (100.0) | 1 ( 33.3) | 1 ( 33.3) | 4 ( 50.0) | 5 ( 45.5) | 10 ( 55.6)
Age (years) |  |  |  |  |  |
n | 4 | 3 | 3 | 8 | 11 | 18
Mean | 67.8 | 44.7 | 60.0 | 42.9 | 47.5 | 51.6
Standard deviation | 6.80 | 23.86 | 20.88 | 13.15 | 16.50 | 17.77
Median | 66.5 | 52.0 | 70.0 | 44.0 | 45.0 | 51.5
Minimum | 61 | 18 | 36 | 22 | 22 | 18
Maximum | 77 | 64 | 74 | 66 | 74 | 77
Age group, n (%) |  |  |  |  |  |
Adult | 4 (100.0) | 3 (100.0) | 3 (100.0) | 8 (100.0) | 11 (100.0) | 18 (100.0)
Pediatric | 0 ( 0.0) | 0 ( 0.0) | 0 ( 0.0) | 0 ( 0.0) | 0 ( 0.0) | 0 ( 0.0)
Sex, n (%) |  |  |  |  |  |
Male | 2 ( 50.0) | 1 ( 33.3) | 1 ( 33.3) | 7 ( 87.5) | 8 ( 72.7) | 11 ( 61.1)
Female | 2 ( 50.0) | 2 ( 66.7) | 2 ( 66.7) | 1 ( 12.5) | 3 ( 27.3) | 7 ( 38.9)
"""

# ==========================================================
# 2️⃣ TABLE PARSER
# ==========================================================

def parse_pipe_table(raw_text):
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
    lines = [l for l in lines if not re.match(r"^\|?-+\|", l)]

    data = []
    max_cols = 0

    for line in lines:
        parts = [c.strip() for c in line.split("|")]
        parts = [p for p in parts if p != ""]
        data.append(parts)
        max_cols = max(max_cols, len(parts))

    for row in data:
        while len(row) < max_cols:
            row.append("")

    return data


# ==========================================================
# 3️⃣ TABLE FORMAT HELPERS
# ==========================================================

def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


# ==========================================================
# 4️⃣ RENDER TABLE INTO TEMPLATE
# ==========================================================
from docx.oxml.ns import qn

def apply_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)


def render_table(template_path, output_path, raw_table_text):
    doc = Document(template_path)

    table_data = parse_pipe_table(raw_table_text)

    for paragraph in doc.paragraphs:
        if "{{ table }}" in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            rows = len(table_data)
            cols = len(table_data[0])

            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            apply_table_borders(table)


            for r_idx, row_data in enumerate(table_data):

                is_section_row = all(cell == "" for cell in row_data[1:])

                for c_idx, value in enumerate(row_data):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = value

                    if c_idx > 0:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if is_section_row:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True

                prevent_row_split(table.rows[r_idx])

            repeat_header(table.rows[0])

            tbl_element = table._element
            parent.insert(index, tbl_element)
            parent.remove(paragraph._element)

            break

    doc.save(output_path)


# ==========================================================
# 5️⃣ MAIN TEST EXECUTION
# ==========================================================

if __name__ == "__main__":
    TEMPLATE_PATH = "CSR.docx" # Must contain {{ table }}
    OUTPUT_PATH = "output.docx"

    render_table(TEMPLATE_PATH, OUTPUT_PATH, RAW_TABLE)

    print("✅ Table rendered successfully into output.docx")